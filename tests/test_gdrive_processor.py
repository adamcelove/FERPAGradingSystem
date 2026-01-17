"""Integration tests for DriveProcessor.

Tests the DriveProcessor orchestrator with mocked Drive API but real pipeline components.
"""

from datetime import datetime
from io import BytesIO
from unittest.mock import MagicMock, patch

import pytest

from ferpa_feedback.gdrive.config import DriveConfig
from ferpa_feedback.gdrive.discovery import DriveDocument, FolderMap, FolderNode
from ferpa_feedback.gdrive.errors import DownloadError
from ferpa_feedback.gdrive.processor import DriveProcessor, ProcessingSummary

# -----------------------------------------------------------------------------
# Fixtures
# -----------------------------------------------------------------------------


@pytest.fixture
def mock_authenticator() -> MagicMock:
    """Create a mock authenticator."""
    auth = MagicMock()
    auth.get_service.return_value = MagicMock()
    auth.service_account_email = "test@example.com"
    return auth


@pytest.fixture
def mock_pipeline() -> MagicMock:
    """Create a mock pipeline."""
    pipeline = MagicMock()
    # Mock process_document to return a document with some data
    mock_doc = MagicMock()
    mock_doc.comments = [MagicMock(grammar_issues=[], anonymization_mappings=[])]
    pipeline.process_document.return_value = mock_doc
    return pipeline


@pytest.fixture
def drive_config() -> DriveConfig:
    """Create a test DriveConfig."""
    return DriveConfig()


@pytest.fixture
def sample_folder_tree() -> FolderNode:
    """Create a sample folder tree for testing.

    Structure:
    Root/
    ├── House1/
    │   └── TeacherA/
    │       ├── September Comments/ (1 doc)
    │       └── Interim 1 Comments/ (1 doc)
    └── House2/
        └── TeacherB/
            └── September Comments/ (1 doc)
    """
    # Create documents
    doc1 = DriveDocument(
        id="doc1",
        name="Comments.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        parent_folder_id="sept1",
        modified_time="2026-01-15T10:00:00Z",
    )
    doc2 = DriveDocument(
        id="doc2",
        name="Comments.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        parent_folder_id="interim1",
        modified_time="2026-01-15T11:00:00Z",
    )
    doc3 = DriveDocument(
        id="doc3",
        name="Comments.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        parent_folder_id="sept2",
        modified_time="2026-01-15T12:00:00Z",
    )

    # Create leaf folders
    sept1 = FolderNode(id="sept1", name="September Comments", parent_id="teachera", depth=3, documents=[doc1])
    sept1.set_path_components(["Root", "House1", "TeacherA", "September Comments"])

    interim1 = FolderNode(id="interim1", name="Interim 1 Comments", parent_id="teachera", depth=3, documents=[doc2])
    interim1.set_path_components(["Root", "House1", "TeacherA", "Interim 1 Comments"])

    sept2 = FolderNode(id="sept2", name="September Comments", parent_id="teacherb", depth=3, documents=[doc3])
    sept2.set_path_components(["Root", "House2", "TeacherB", "September Comments"])

    # Create teacher folders
    teachera = FolderNode(id="teachera", name="TeacherA", parent_id="house1", depth=2, children=[sept1, interim1])
    teachera.set_path_components(["Root", "House1", "TeacherA"])

    teacherb = FolderNode(id="teacherb", name="TeacherB", parent_id="house2", depth=2, children=[sept2])
    teacherb.set_path_components(["Root", "House2", "TeacherB"])

    # Create house folders
    house1 = FolderNode(id="house1", name="House1", parent_id="root", depth=1, children=[teachera])
    house1.set_path_components(["Root", "House1"])

    house2 = FolderNode(id="house2", name="House2", parent_id="root", depth=1, children=[teacherb])
    house2.set_path_components(["Root", "House2"])

    # Create root
    root = FolderNode(id="root", name="Root", parent_id=None, depth=0, children=[house1, house2])
    root.set_path_components(["Root"])

    return root


@pytest.fixture
def sample_folder_map(sample_folder_tree: FolderNode) -> FolderMap:
    """Create a FolderMap from the sample tree."""
    return FolderMap(
        root=sample_folder_tree,
        discovered_at=datetime(2026, 1, 17, 10, 0, 0),
        total_folders=7,
        total_documents=3,
    )


@pytest.fixture
def minimal_docx_bytes() -> bytes:
    """Create minimal valid .docx bytes."""
    from docx import Document

    doc = Document()
    header = doc.add_paragraph()
    header.add_run("Student Name: Test Student\tGrade: A").bold = True
    doc.add_paragraph("This is a test comment.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# -----------------------------------------------------------------------------
# Test DriveProcessor initialization
# -----------------------------------------------------------------------------


class TestDriveProcessorInit:
    """Tests for DriveProcessor initialization."""

    def test_processor_initializes_with_required_args(
        self, mock_authenticator: MagicMock, mock_pipeline: MagicMock
    ) -> None:
        """Processor initializes with authenticator and pipeline."""
        processor = DriveProcessor(mock_authenticator, mock_pipeline)

        assert processor is not None

    def test_processor_initializes_with_config(
        self, mock_authenticator: MagicMock, mock_pipeline: MagicMock, drive_config: DriveConfig
    ) -> None:
        """Processor accepts optional config."""
        processor = DriveProcessor(mock_authenticator, mock_pipeline, drive_config)

        assert processor is not None


# -----------------------------------------------------------------------------
# Test list_folders
# -----------------------------------------------------------------------------


class TestListFolders:
    """Tests for DriveProcessor.list_folders method."""

    def test_list_folders_returns_folder_map(
        self,
        mock_authenticator: MagicMock,
        mock_pipeline: MagicMock,
        sample_folder_map: FolderMap,
    ) -> None:
        """list_folders returns a FolderMap."""
        processor = DriveProcessor(mock_authenticator, mock_pipeline)

        with patch.object(processor, "_discovery") as mock_discovery:
            mock_discovery.discover_structure.return_value = sample_folder_map

            result = processor.list_folders("root_id")

        assert isinstance(result, FolderMap)
        assert result.total_folders == 7
        assert result.total_documents == 3


# -----------------------------------------------------------------------------
# Test process with pattern filtering
# -----------------------------------------------------------------------------


class TestProcessorFiltering:
    """Tests for pattern-based folder filtering."""

    def test_processor_filters_by_pattern(
        self,
        mock_authenticator: MagicMock,
        mock_pipeline: MagicMock,
        sample_folder_map: FolderMap,
        minimal_docx_bytes: bytes,
    ) -> None:
        """Processor filters folders by pattern."""
        processor = DriveProcessor(mock_authenticator, mock_pipeline)

        # Mock discovery
        with patch.object(processor, "_discovery") as mock_discovery:
            mock_discovery.discover_structure.return_value = sample_folder_map

            # Mock downloader
            with patch.object(processor, "_downloader") as mock_downloader:
                # Return downloaded documents for September* pattern (2 docs)
                mock_downloaded = MagicMock()
                mock_downloaded.drive_document = DriveDocument(
                    id="doc1", name="Comments.docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    parent_folder_id="sept1", modified_time="2026-01-15T10:00:00Z",
                )
                mock_downloaded.content = BytesIO(minimal_docx_bytes)
                mock_downloader.download_batch.return_value = iter([mock_downloaded, mock_downloaded])

                # Mock uploader
                with patch.object(processor, "_uploader") as mock_uploader:
                    mock_uploader.ensure_output_folder.return_value = "output_folder"
                    mock_upload_result = MagicMock()
                    mock_upload_result.success = True
                    mock_uploader.upload_grammar_report.return_value = mock_upload_result
                    mock_uploader.upload_anonymized_output.return_value = mock_upload_result

                    summary = processor.process(
                        root_folder_id="root_id",
                        target_patterns=["September*"],
                    )

        # Should have filtered to September folders only (2 docs)
        assert summary.total_documents == 2

    def test_processor_processes_all_without_pattern(
        self,
        mock_authenticator: MagicMock,
        mock_pipeline: MagicMock,
        sample_folder_map: FolderMap,
        minimal_docx_bytes: bytes,
    ) -> None:
        """Processor processes all leaf folders when no pattern specified."""
        processor = DriveProcessor(mock_authenticator, mock_pipeline)

        with patch.object(processor, "_discovery") as mock_discovery:
            mock_discovery.discover_structure.return_value = sample_folder_map

            with patch.object(processor, "_downloader") as mock_downloader:
                mock_downloaded = MagicMock()
                mock_downloaded.drive_document = DriveDocument(
                    id="doc1", name="Comments.docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    parent_folder_id="sept1", modified_time="2026-01-15T10:00:00Z",
                )
                mock_downloaded.content = BytesIO(minimal_docx_bytes)
                # Return 3 docs for all leaf folders
                mock_downloader.download_batch.return_value = iter([
                    mock_downloaded, mock_downloaded, mock_downloaded
                ])

                with patch.object(processor, "_uploader") as mock_uploader:
                    mock_uploader.ensure_output_folder.return_value = "output_folder"
                    mock_upload_result = MagicMock()
                    mock_upload_result.success = True
                    mock_uploader.upload_grammar_report.return_value = mock_upload_result
                    mock_uploader.upload_anonymized_output.return_value = mock_upload_result

                    summary = processor.process(
                        root_folder_id="root_id",
                        target_patterns=None,  # No filter
                    )

        # Should process all 3 documents
        assert summary.total_documents == 3


# -----------------------------------------------------------------------------
# Test error handling
# -----------------------------------------------------------------------------


class TestProcessorErrorHandling:
    """Tests for error handling and continue-on-error behavior."""

    def test_processor_continues_on_document_error(
        self,
        mock_authenticator: MagicMock,
        mock_pipeline: MagicMock,
        sample_folder_map: FolderMap,
        minimal_docx_bytes: bytes,
    ) -> None:
        """Processor continues processing when individual documents fail."""
        processor = DriveProcessor(mock_authenticator, mock_pipeline)

        with patch.object(processor, "_discovery") as mock_discovery:
            mock_discovery.discover_structure.return_value = sample_folder_map

            with patch.object(processor, "_downloader") as mock_downloader:
                # Create successful download result
                mock_success = MagicMock()
                mock_success.drive_document = DriveDocument(
                    id="doc1", name="Good.docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    parent_folder_id="sept1", modified_time="2026-01-15T10:00:00Z",
                )
                mock_success.content = BytesIO(minimal_docx_bytes)

                # Set up download_document to succeed twice and fail once
                # Documents are: doc1 (sept1), doc2 (interim1), doc3 (sept2)
                call_count = [0]

                def download_side_effect(doc: DriveDocument) -> MagicMock:
                    call_count[0] += 1
                    if call_count[0] == 2:  # Second document fails
                        raise DownloadError("Failed to download", file_id=doc.id)
                    # Return fresh BytesIO for each successful call
                    result = MagicMock()
                    result.drive_document = doc
                    result.content = BytesIO(minimal_docx_bytes)
                    return result

                mock_downloader.download_document.side_effect = download_side_effect

                with patch.object(processor, "_uploader") as mock_uploader:
                    mock_uploader.ensure_output_folder.return_value = "output_folder"
                    mock_upload_result = MagicMock()
                    mock_upload_result.success = True
                    mock_uploader.upload_grammar_report.return_value = mock_upload_result
                    mock_uploader.upload_anonymized_output.return_value = mock_upload_result

                    summary = processor.process(root_folder_id="root_id")

        # Should have processed 3 total, 2 successful, 1 failed
        assert summary.total_documents == 3
        assert summary.successful == 2
        assert summary.failed == 1
        assert len(summary.errors) == 1


# -----------------------------------------------------------------------------
# Test summary generation
# -----------------------------------------------------------------------------


class TestProcessorSummary:
    """Tests for ProcessingSummary generation."""

    def test_processor_generates_summary(
        self,
        mock_authenticator: MagicMock,
        mock_pipeline: MagicMock,
        sample_folder_map: FolderMap,
        minimal_docx_bytes: bytes,
    ) -> None:
        """Processor generates complete summary."""
        # Configure mock pipeline to return document with grammar issues
        mock_doc = MagicMock()
        mock_comment = MagicMock()
        mock_comment.grammar_issues = [MagicMock(), MagicMock()]  # 2 issues
        mock_comment.anonymization_mappings = [MagicMock()]  # 1 PII
        mock_doc.comments = [mock_comment]
        mock_pipeline.process_document.return_value = mock_doc

        processor = DriveProcessor(mock_authenticator, mock_pipeline)

        with patch.object(processor, "_discovery") as mock_discovery:
            mock_discovery.discover_structure.return_value = sample_folder_map

            with patch.object(processor, "_downloader") as mock_downloader:
                mock_downloaded = MagicMock()
                mock_downloaded.drive_document = DriveDocument(
                    id="doc1", name="Comments.docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    parent_folder_id="sept1", modified_time="2026-01-15T10:00:00Z",
                )
                mock_downloaded.content = BytesIO(minimal_docx_bytes)
                mock_downloader.download_batch.return_value = iter([mock_downloaded])

                with patch.object(processor, "_uploader") as mock_uploader:
                    mock_uploader.ensure_output_folder.return_value = "output_folder"
                    mock_upload_result = MagicMock()
                    mock_upload_result.success = True
                    mock_uploader.upload_grammar_report.return_value = mock_upload_result
                    mock_uploader.upload_anonymized_output.return_value = mock_upload_result

                    summary = processor.process(
                        root_folder_id="root_id",
                        target_patterns=["September*"],
                    )

        assert isinstance(summary, ProcessingSummary)
        assert summary.started_at is not None
        assert summary.completed_at is not None
        assert summary.duration_seconds >= 0
        assert summary.folder_map is not None
        assert summary.successful >= 0
        assert summary.failed >= 0

    def test_summary_duration_calculation(self) -> None:
        """ProcessingSummary.duration_seconds is calculated correctly."""
        start = datetime(2026, 1, 17, 10, 0, 0)
        end = datetime(2026, 1, 17, 10, 1, 30)  # 90 seconds later

        summary = ProcessingSummary(
            started_at=start,
            completed_at=end,
            folder_map=MagicMock(),
            target_folders=[],
            total_documents=10,
            successful=8,
            failed=2,
            grammar_issues_found=5,
            pii_instances_replaced=3,
            uploads_completed=16,
        )

        assert summary.duration_seconds == 90.0

    def test_summary_success_rate(self) -> None:
        """ProcessingSummary.success_rate is calculated correctly."""
        summary = ProcessingSummary(
            started_at=datetime.now(),
            completed_at=datetime.now(),
            folder_map=MagicMock(),
            target_folders=[],
            total_documents=10,
            successful=8,
            failed=2,
            grammar_issues_found=0,
            pii_instances_replaced=0,
            uploads_completed=0,
        )

        assert summary.success_rate == 80.0

    def test_summary_success_rate_zero_documents(self) -> None:
        """ProcessingSummary.success_rate handles zero documents."""
        summary = ProcessingSummary(
            started_at=datetime.now(),
            completed_at=datetime.now(),
            folder_map=MagicMock(),
            target_folders=[],
            total_documents=0,
            successful=0,
            failed=0,
            grammar_issues_found=0,
            pii_instances_replaced=0,
            uploads_completed=0,
        )

        assert summary.success_rate == 0.0


# -----------------------------------------------------------------------------
# Test dry run mode
# -----------------------------------------------------------------------------


class TestDryRunMode:
    """Tests for dry run mode."""

    def test_dry_run_does_not_process(
        self,
        mock_authenticator: MagicMock,
        mock_pipeline: MagicMock,
        sample_folder_map: FolderMap,
    ) -> None:
        """Dry run mode discovers folders but doesn't process."""
        processor = DriveProcessor(mock_authenticator, mock_pipeline)

        with patch.object(processor, "_discovery") as mock_discovery:
            mock_discovery.discover_structure.return_value = sample_folder_map

            summary = processor.process(
                root_folder_id="root_id",
                dry_run=True,
            )

        # Should have found documents but not processed them
        assert summary.total_documents == 3
        assert summary.successful == 0
        assert summary.failed == 0
        # Pipeline should not have been called
        mock_pipeline.process_document.assert_not_called()
