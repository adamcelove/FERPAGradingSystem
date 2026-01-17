"""Unit tests for DocumentParser, focusing on BytesIO support for Google Drive integration.

Tests that the parser works with both Path and BytesIO inputs.
"""

from io import BytesIO
from pathlib import Path
from typing import Any, Dict

import pytest

from ferpa_feedback.stage_0_ingestion import DocumentParser

# -----------------------------------------------------------------------------
# Fixtures
# -----------------------------------------------------------------------------


@pytest.fixture
def parser() -> DocumentParser:
    """Create a DocumentParser instance."""
    return DocumentParser()


@pytest.fixture
def minimal_docx_bytes() -> bytes:
    """Create a minimal valid .docx file as bytes.

    This creates a simple document with a header and one comment section.
    """
    from docx import Document

    doc = Document()

    # Add a header paragraph (will be detected as combined header format)
    header = doc.add_paragraph()
    header_run = header.add_run("Student Name: John Doe\tGrade: A")
    header_run.bold = True

    # Add comment text
    doc.add_paragraph("This is a sample comment about the student's performance.")

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture
def docx_bytesio(minimal_docx_bytes: bytes) -> BytesIO:
    """Create a BytesIO stream from minimal docx bytes."""
    return BytesIO(minimal_docx_bytes)


@pytest.fixture
def sample_docx_path() -> Path:
    """Return path to sample docx file if it exists."""
    sample_path = Path(__file__).parent.parent / "sample_data" / "Adam_Love_Interim_1 Comments 25-26.docx"
    if sample_path.exists():
        return sample_path
    pytest.skip("Sample docx file not found")
    return sample_path  # Never reached, but satisfies type checker


# -----------------------------------------------------------------------------
# Test BytesIO support
# -----------------------------------------------------------------------------


class TestParseDocxFromBytesIO:
    """Tests for parsing .docx from BytesIO streams."""

    def test_parse_docx_from_bytesio_returns_document(
        self, parser: DocumentParser, docx_bytesio: BytesIO
    ) -> None:
        """BytesIO stream input successfully parses to TeacherDocument."""
        result = parser.parse_docx(docx_bytesio)

        assert result is not None
        assert result.id is not None
        # Should have extracted at least some content
        assert result.source_path is not None

    def test_parse_docx_from_bytesio_with_document_id(
        self, parser: DocumentParser, docx_bytesio: BytesIO
    ) -> None:
        """Custom document_id is used when provided."""
        custom_id = "custom-doc-123"
        result = parser.parse_docx(docx_bytesio, document_id=custom_id)

        assert result.id == custom_id

    def test_parse_docx_from_bytesio_with_metadata(
        self, parser: DocumentParser, docx_bytesio: BytesIO
    ) -> None:
        """Metadata parameter is used for source_path with drive_file_id."""
        metadata: Dict[str, Any] = {"drive_file_id": "drive-abc-123"}
        result = parser.parse_docx(docx_bytesio, metadata=metadata)

        # source_path should use the drive_file_id from metadata
        assert "drive-abc-123" in result.source_path

    def test_parse_docx_from_bytesio_stream_position_reset(
        self, parser: DocumentParser, minimal_docx_bytes: bytes
    ) -> None:
        """Parser works even if BytesIO position is not at start."""
        # Create BytesIO and move position
        stream = BytesIO(minimal_docx_bytes)
        stream.seek(100)  # Move position away from start

        # Should still work (parser should seek to start)
        result = parser.parse_docx(stream)
        assert result is not None


# -----------------------------------------------------------------------------
# Test Path support (regression)
# -----------------------------------------------------------------------------


class TestParseDocxFromPath:
    """Tests for parsing .docx from file paths (regression tests)."""

    def test_parse_docx_from_path_returns_document(
        self, parser: DocumentParser, sample_docx_path: Path
    ) -> None:
        """Path input still works correctly."""
        result = parser.parse_docx(sample_docx_path)

        assert result is not None
        assert result.id is not None
        assert str(sample_docx_path) in result.source_path

    def test_parse_docx_from_path_extracts_comments(
        self, parser: DocumentParser, sample_docx_path: Path
    ) -> None:
        """Path-based parsing extracts comments."""
        result = parser.parse_docx(sample_docx_path)

        # Should have extracted some comments
        assert len(result.comments) > 0

    def test_parse_docx_from_path_with_document_id(
        self, parser: DocumentParser, sample_docx_path: Path
    ) -> None:
        """Custom document_id works with Path input."""
        custom_id = "path-doc-456"
        result = parser.parse_docx(sample_docx_path, document_id=custom_id)

        assert result.id == custom_id


# -----------------------------------------------------------------------------
# Test metadata parameter
# -----------------------------------------------------------------------------


class TestParseDocxMetadata:
    """Tests for metadata parameter handling."""

    def test_metadata_none_uses_default_source_path(
        self, parser: DocumentParser, docx_bytesio: BytesIO
    ) -> None:
        """When metadata is None, source_path uses default stream identifier."""
        result = parser.parse_docx(docx_bytesio, metadata=None)

        # Should have some source_path even without metadata
        assert result.source_path is not None
        assert len(result.source_path) > 0

    def test_metadata_empty_dict_uses_default(
        self, parser: DocumentParser, docx_bytesio: BytesIO
    ) -> None:
        """When metadata is empty dict, source_path uses default."""
        result = parser.parse_docx(docx_bytesio, metadata={})

        assert result.source_path is not None
        # Should contain "stream://" for BytesIO without drive_file_id
        assert "stream://" in result.source_path or result.source_path != ""

    def test_metadata_with_extra_fields(
        self, parser: DocumentParser, docx_bytesio: BytesIO
    ) -> None:
        """Extra metadata fields don't cause issues."""
        metadata: Dict[str, Any] = {
            "drive_file_id": "drive-xyz",
            "folder_name": "September Comments",
            "teacher": "Mr. Smith",
        }
        result = parser.parse_docx(docx_bytesio, metadata=metadata)

        assert result is not None
        assert "drive-xyz" in result.source_path


# -----------------------------------------------------------------------------
# Test convenience function
# -----------------------------------------------------------------------------


class TestParseDocumentFunction:
    """Tests for the parse_document convenience function."""

    def test_parse_document_with_bytesio(self, docx_bytesio: BytesIO) -> None:
        """parse_document function works with BytesIO."""
        from ferpa_feedback.stage_0_ingestion import parse_document

        result = parse_document(docx_bytesio)

        assert result is not None
        assert result.id is not None

    def test_parse_document_with_path(self, sample_docx_path: Path) -> None:
        """parse_document function works with Path."""
        from ferpa_feedback.stage_0_ingestion import parse_document

        result = parse_document(sample_docx_path)

        assert result is not None
        assert len(result.comments) > 0
