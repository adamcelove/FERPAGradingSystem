"""
Stage 0: Document Ingestion

Parses teacher comment documents (Word docs) into structured StudentComment records.

Key improvements over original:
- Format auto-detection (combined header vs separate lines vs table)
- Preserves full_name and first_name separately for anonymization
- Validation flags for malformed or unexpected comments
- Handles dash variants (-, -, -) and inconsistent spacing
- More robust name patterns (handles O'Brien, McDonald, hyphenated names)

This stage is 100% local - no external API calls.
"""

from __future__ import annotations

import re
import uuid
from collections.abc import Iterator
from enum import Enum, auto
from pathlib import Path
from typing import Any

import structlog
from docx import Document as DocxDocument
from docx.document import Document
from docx.table import Table

from ferpa_feedback.models import StudentComment, TeacherDocument

logger = structlog.get_logger()


# =============================================================================
# Enums for parsing and validation
# =============================================================================

class DocumentFormat(Enum):
    """Detected document format types."""
    COMBINED_HEADER = auto()    # "LastName, FirstName - Grade" on one line
    SEPARATE_HEADER = auto()    # Name on one line, grade on next
    TABLE = auto()              # Tabular format
    UNKNOWN = auto()


class ValidationFlag(Enum):
    """Flags for potential issues with parsed comments."""
    SHORT_COMMENT = auto()          # Comment under expected word count
    LONG_COMMENT = auto()           # Comment over expected word count
    NO_GRADE = auto()               # Grade not detected
    UNUSUAL_GRADE = auto()          # Grade format unexpected
    NAME_IN_COMMENT_MISMATCH = auto()  # First name in comment doesn't match header
    NO_FIRST_NAME_USAGE = auto()    # Comment doesn't use student's first name
    POSSIBLE_WRONG_NAME = auto()    # Different name appears frequently in comment
    EMPTY_COMMENT = auto()          # No comment text found
    PARSING_UNCERTAIN = auto()      # Parser confidence low


# =============================================================================
# Parser
# =============================================================================

class DocumentParser:
    """
    Parses teacher comment documents into structured records.

    Supports multiple document formats with auto-detection:
    1. Combined header: "LastName, FirstName - Grade" (most common)
    2. Separate header: Name line, then grade line, then comment
    3. Table format: Columns for name, grade, comment
    """

    # Expected comment length (from guidelines: 200-250 words)
    MIN_WORD_COUNT = 150   # Allow some buffer below guideline
    MAX_WORD_COUNT = 300   # Allow some buffer above guideline
    TARGET_MIN = 200
    TARGET_MAX = 250

    # Combined header pattern: "LastName, FirstName - Grade"
    # Handles: dash variants, inconsistent spacing, various name formats, Unicode
    COMBINED_HEADER_PATTERN = re.compile(
        r"^([\w][\w'\-]+),\s*"               # Last name (Unicode word chars, apostrophe, hyphen)
        r"([\w][\w'\- ]*?)\s*"               # First name (may have middle, spaces)
        r"[-\u2013\u2014]\s*"                # Dash separator (any variant)
        r"([A-F][+-]?|\d{1,3}(?:\.\d+)?%?)$", # Grade
        re.IGNORECASE | re.UNICODE
    )

    # Separate name-only patterns (no grade on line)
    NAME_ONLY_PATTERNS = [
        re.compile(r"^([A-Za-z][A-Za-z'\-]+),\s*([A-Za-z][A-Za-z'\- ]+)$"),  # "Last, First"
        re.compile(r"^([A-Za-z][A-Za-z'\- ]+)\s+([A-Za-z][A-Za-z'\-]+)$"),   # "First Last"
    ]

    # Grade-only patterns
    GRADE_PATTERNS = [
        re.compile(r"^Grade[:\s]*([A-F][+-]?)$", re.IGNORECASE),
        re.compile(r"^([A-F][+-]?)$"),
        re.compile(r"^Grade[:\s]*(\d{1,3}(?:\.\d+)?%?)$", re.IGNORECASE),
        re.compile(r"^(\d{1,3}(?:\.\d+)?%?)$"),
    ]

    # Section delimiters
    DELIMITER_PATTERNS = [
        re.compile(r"^[-=_]{3,}$"),
        re.compile(r"^\*{3,}$"),
    ]

    def __init__(self) -> None:
        self._format_detection_cache: dict[str, DocumentFormat] = {}

    def parse_docx(self, file_path: Path, document_id: str | None = None) -> TeacherDocument:
        """
        Parse a Word document into a TeacherDocument.

        Args:
            file_path: Path to the .docx file
            document_id: Optional ID (generated if not provided)

        Returns:
            TeacherDocument with extracted comments
        """
        document_id = document_id or str(uuid.uuid4())
        doc: Document = DocxDocument(str(file_path))

        logger.info("parsing_document", path=str(file_path), doc_id=document_id)

        # Detect format
        detected_format = self._detect_format(doc)
        logger.info("format_detected", format=detected_format.name)

        warnings = []

        # Parse based on detected format
        if detected_format == DocumentFormat.TABLE:
            if doc.tables:
                comments = list(self._parse_table_format(doc.tables[0], document_id))
            else:
                warnings.append("Table format detected but no tables found")
                comments = []
        elif detected_format == DocumentFormat.COMBINED_HEADER:
            comments = list(self._parse_combined_header_format(doc, document_id))
        elif detected_format == DocumentFormat.SEPARATE_HEADER:
            comments = list(self._parse_separate_header_format(doc, document_id))
        else:
            warnings.append("Could not detect document format, attempting combined header parse")
            comments = list(self._parse_combined_header_format(doc, document_id))

        logger.info(
            "parsing_complete",
            comment_count=len(comments),
        )

        return TeacherDocument(
            id=document_id,
            source_path=str(file_path),
            teacher_name="",  # To be filled from filename or external source
            class_name="",    # To be filled from filename or external source
            term="",          # To be filled from filename or external source
            comments=comments,
        )

    def _detect_format(self, doc: Any) -> DocumentFormat:
        """
        Auto-detect the document format by examining structure.
        """
        # Check for tables first
        if doc.tables and len(doc.tables[0].rows) > 1:
            # Verify it looks like a comment table
            header_text = " ".join(c.text.lower() for c in doc.tables[0].rows[0].cells)
            if any(kw in header_text for kw in ["name", "student", "comment", "grade"]):
                return DocumentFormat.TABLE

        # Examine first several non-empty paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()][:10]

        combined_matches = 0
        name_only_matches = 0

        for para in paragraphs:
            if self.COMBINED_HEADER_PATTERN.match(para):
                combined_matches += 1
            elif any(p.match(para) for p in self.NAME_ONLY_PATTERNS):
                name_only_matches += 1

        if combined_matches >= 1:
            return DocumentFormat.COMBINED_HEADER
        elif name_only_matches >= 1:
            return DocumentFormat.SEPARATE_HEADER

        return DocumentFormat.UNKNOWN

    def _parse_combined_header_format(
        self, doc: Any, document_id: str
    ) -> Iterator[StudentComment]:
        """
        Parse documents with combined "Name - Grade" headers.

        Expected format:
            LastName, FirstName - A+

            Comment paragraph text...

            NextLastName, NextFirstName - B

            Next comment...
        """
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        current_header = None
        current_last = None
        current_first = None
        current_grade = None
        comment_lines: list[str] = []
        section_index = 0

        for para in paragraphs:
            match = self.COMBINED_HEADER_PATTERN.match(para)

            if match:
                # Save previous student if exists
                if current_header is not None and comment_lines:
                    section_index += 1
                    yield self._create_comment(
                        document_id=document_id,
                        section_index=section_index,
                        last_name=current_last or "",
                        first_name=current_first or "",
                        grade=current_grade,
                        comment_lines=comment_lines,
                    )

                # Start new student
                current_header = para
                current_last = match.group(1).strip()
                current_first = match.group(2).strip()
                current_grade = match.group(3).strip()
                comment_lines = []
            else:
                # Accumulate comment text
                comment_lines.append(para)

        # Don't forget last student
        if current_header is not None and comment_lines:
            section_index += 1
            yield self._create_comment(
                document_id=document_id,
                section_index=section_index,
                last_name=current_last or "",
                first_name=current_first or "",
                grade=current_grade,
                comment_lines=comment_lines,
            )

    def _parse_separate_header_format(
        self, doc: Any, document_id: str
    ) -> Iterator[StudentComment]:
        """
        Parse documents with name and grade on separate lines.
        """
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        current_header = None
        current_last = None
        current_first = None
        current_grade = None
        comment_lines: list[str] = []
        section_index = 0

        for para in paragraphs:
            # Check for delimiter
            if any(p.match(para) for p in self.DELIMITER_PATTERNS):
                if current_header and comment_lines:
                    section_index += 1
                    yield self._create_comment(
                        document_id=document_id,
                        section_index=section_index,
                        last_name=current_last or "",
                        first_name=current_first or "",
                        grade=current_grade,
                        comment_lines=comment_lines,
                    )
                current_header = None
                current_last = None
                current_first = None
                current_grade = None
                comment_lines = []
                continue

            # Check for name-only line
            name_match = None
            for pattern in self.NAME_ONLY_PATTERNS:
                name_match = pattern.match(para)
                if name_match:
                    break

            if name_match:
                # Save previous if exists
                if current_header and comment_lines:
                    section_index += 1
                    yield self._create_comment(
                        document_id=document_id,
                        section_index=section_index,
                        last_name=current_last or "",
                        first_name=current_first or "",
                        grade=current_grade,
                        comment_lines=comment_lines,
                    )

                current_header = para
                # Determine if Last, First or First Last
                if "," in para:
                    current_last = name_match.group(1).strip()
                    current_first = name_match.group(2).strip()
                else:
                    current_first = name_match.group(1).strip()
                    current_last = name_match.group(2).strip()
                current_grade = None
                comment_lines = []
                continue

            # Check for grade line
            grade_match = None
            for pattern in self.GRADE_PATTERNS:
                grade_match = pattern.match(para)
                if grade_match:
                    break

            if grade_match and current_header:
                current_grade = grade_match.group(1).strip()
                continue

            # Otherwise it's comment text
            if current_header:
                comment_lines.append(para)

        # Last student
        if current_header and comment_lines:
            section_index += 1
            yield self._create_comment(
                document_id=document_id,
                section_index=section_index,
                last_name=current_last or "",
                first_name=current_first or "",
                grade=current_grade,
                comment_lines=comment_lines,
            )

    def _parse_table_format(
        self, table: Table, document_id: str
    ) -> Iterator[StudentComment]:
        """Parse a table-formatted document."""
        header_row = [cell.text.strip().lower() for cell in table.rows[0].cells]

        # Find columns
        name_col = self._find_column(header_row, ["name", "student", "student name"])
        grade_col = self._find_column(header_row, ["grade", "mark", "score"])
        comment_col = self._find_column(header_row, ["comment", "comments", "feedback", "notes"])

        # Try to find separate first/last columns
        first_col = self._find_column(header_row, ["first", "first name", "given"])
        last_col = self._find_column(header_row, ["last", "last name", "surname", "family"])

        if comment_col is None:
            logger.warning("table_missing_comment_column", headers=header_row)
            return

        for idx, row in enumerate(table.rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]

            # Extract name
            if first_col is not None and last_col is not None:
                first_name = cells[first_col] if first_col < len(cells) else ""
                last_name = cells[last_col] if last_col < len(cells) else ""
            elif name_col is not None:
                raw_name = cells[name_col] if name_col < len(cells) else ""
                last_name, first_name = self._parse_name_string(raw_name)
            else:
                continue

            grade = cells[grade_col] if grade_col is not None and grade_col < len(cells) else ""
            comment = cells[comment_col] if comment_col < len(cells) else ""

            if not comment:
                continue

            yield self._create_comment(
                document_id=document_id,
                section_index=idx,
                last_name=last_name,
                first_name=first_name,
                grade=grade,
                comment_lines=[comment],
            )

    def _create_comment(
        self,
        document_id: str,
        section_index: int,
        last_name: str,
        first_name: str,
        grade: str | None,
        comment_lines: list[str],
    ) -> StudentComment:
        """Create a Pydantic StudentComment with computed fields."""
        comment_text = " ".join(comment_lines)

        # Build student name from parts (using "LastName, FirstName" format)
        student_name = f"{last_name}, {first_name}" if last_name and first_name else (last_name or first_name)

        return StudentComment(
            id=f"{document_id}-{section_index}",
            document_id=document_id,
            section_index=section_index,
            student_name=student_name,
            grade=grade or "",
            comment_text=comment_text,
            # Analysis fields left as defaults (None/empty) - populated by later stages
        )

    def _find_column(self, headers: list[str], candidates: list[str]) -> int | None:
        """Find the index of a column matching any candidate name."""
        for idx, header in enumerate(headers):
            if any(c in header for c in candidates):
                return idx
        return None

    def _parse_name_string(self, name_str: str) -> tuple[str, str]:
        """Parse a name string into (last_name, first_name)."""
        if "," in name_str:
            parts = name_str.split(",", 1)
            return parts[0].strip(), parts[1].strip()
        else:
            parts = name_str.split()
            if len(parts) >= 2:
                return parts[-1], " ".join(parts[:-1])
            return name_str, ""


# =============================================================================
# Roster Integration
# =============================================================================

class RosterLoader:
    """Loads student rosters and can match against parsed comments."""

    @staticmethod
    def from_csv(file_path: Path) -> list[dict[str, str]]:
        """
        Load roster from CSV file.

        Expected columns: student_id, first_name, last_name, preferred_name (optional)
        """
        import csv

        roster: list[dict[str, str]] = []
        with open(file_path, newline='', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                roster.append({
                    'student_id': row.get('student_id', row.get('id', '')),
                    'first_name': row.get('first_name', row.get('first', '')),
                    'last_name': row.get('last_name', row.get('last', '')),
                    'preferred_name': row.get('preferred_name', row.get('nickname', '')),
                })

        logger.info("roster_loaded", count=len(roster), source=str(file_path))
        return roster

    @staticmethod
    def match_comment_to_roster(
        comment: StudentComment, roster: list[dict[str, str]]
    ) -> dict[str, str] | None:
        """
        Find the roster entry matching a comment's student.

        Returns the matching roster entry or None.
        """
        # Extract last/first from student_name (assumes "Last, First" format)
        if "," in comment.student_name:
            parts = comment.student_name.split(",", 1)
            comment_last = parts[0].strip().lower()
            comment_first = parts[1].strip().lower()
        else:
            # Try splitting by space (First Last)
            parts = comment.student_name.split()
            if len(parts) >= 2:
                comment_last = parts[-1].lower()
                comment_first = " ".join(parts[:-1]).lower()
            else:
                comment_last = comment.student_name.lower()
                comment_first = ""

        for entry in roster:
            roster_last = entry['last_name'].lower()
            roster_first = entry['first_name'].lower()
            roster_pref = entry.get('preferred_name', '').lower()

            if roster_last == comment_last and (roster_first == comment_first or roster_pref == comment_first):
                return entry

        return None


# =============================================================================
# Convenience Functions
# =============================================================================

def parse_document(file_path: str | Path) -> TeacherDocument:
    """
    Parse a document file into a TeacherDocument.

    Args:
        file_path: Path to .docx file

    Returns:
        Parsed TeacherDocument
    """
    parser = DocumentParser()
    return parser.parse_docx(Path(file_path))


def print_validation_report(doc: TeacherDocument) -> None:
    """Print a human-readable validation report."""
    print(f"\n{'='*60}")
    print(f"Document: {doc.source_path}")
    print(f"Total comments: {len(doc.comments)}")
    print(f"Needs review: {doc.needs_review_count}")

    if doc.needs_review_count > 0:
        print("\nComments Needing Review:")
        for c in doc.comments:
            if c.needs_review:
                print(f"\n  [{c.section_index}] {c.student_name} (Grade: {c.grade})")
                for reason in c.review_reasons:
                    print(f"      - {reason}")

    print(f"\n{'='*60}\n")


# =============================================================================
# Main (for testing)
# =============================================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        doc = parse_document(sys.argv[1])
        print_validation_report(doc)

        print("\nParsed Comments:")
        for c in doc.comments:
            print(f"\n--- {c.student_name} ({c.grade}) ---")
            print(f"Comment preview: {c.comment_text[:100]}...")
    else:
        print("Usage: python stage_0_ingestion.py <document.docx>")
