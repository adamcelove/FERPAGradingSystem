"""
Stage 0: Document Ingestion (Improved)

Parses teacher comment documents (Word docs) into structured StudentComment records.

Key improvements over original:
- Format auto-detection (combined header vs separate lines vs table)
- Preserves full_name and first_name separately for anonymization
- Validation flags for malformed or unexpected comments
- Handles dash variants (-, –, —) and inconsistent spacing
- More robust name patterns (handles O'Brien, McDonald, hyphenated names)

This stage is 100% local - no external API calls.
"""

import re
import uuid
from dataclasses import dataclass, field
from enum import Enum, auto
from pathlib import Path
from typing import Iterator

import structlog
from docx import Document
from docx.table import Table

logger = structlog.get_logger()


# =============================================================================
# Models
# =============================================================================

class DocumentFormat(Enum):
    """Detected document format types."""
    COMBINED_HEADER = auto()    # "LastName, FirstName - Grade" on one line
    SEPARATE_HEADER = auto()    # Name on one line, grade on next
    TABLE = auto()              # Tabular format
    UNKNOWN = auto()


class ValidationFlag(Enum):
    """Flags for potential issues with parsed comments."""
    SHORT_COMMENT = auto()          # Comment under expected word count
    LONG_COMMENT = auto()           # Comment over expected word count
    NO_GRADE = auto()               # Grade not detected
    UNUSUAL_GRADE = auto()          # Grade format unexpected
    NAME_IN_COMMENT_MISMATCH = auto()  # First name in comment doesn't match header
    NO_FIRST_NAME_USAGE = auto()    # Comment doesn't use student's first name
    POSSIBLE_WRONG_NAME = auto()    # Different name appears frequently in comment
    EMPTY_COMMENT = auto()          # No comment text found
    PARSING_UNCERTAIN = auto()      # Parser confidence low


@dataclass
class StudentComment:
    """A single student's comment record."""
    # Required fields (no defaults)
    id: str
    document_id: str
    section_index: int
    
    # Name fields - preserved separately for anonymization
    raw_header: str              # Original header text as found
    full_name: str               # "LastName, FirstName" normalized
    last_name: str
    first_name: str
    grade: str
    comment_text: str
    
    # Fields with defaults
    preferred_name: str = ""     # If different from first_name (from roster)
    word_count: int = 0
    
    # Validation
    flags: list[ValidationFlag] = field(default_factory=list)
    flag_details: dict[str, str] = field(default_factory=dict)
    names_found_in_comment: list[str] = field(default_factory=list)


@dataclass
class TeacherDocument:
    """A parsed teacher comment document."""
    id: str
    source_path: str
    
    # Metadata (to be filled from filename or external source)
    teacher_name: str = ""
    class_name: str = ""
    term: str = ""
    
    # Parsing results
    detected_format: DocumentFormat = DocumentFormat.UNKNOWN
    comments: list[StudentComment] = field(default_factory=list)
    
    # Document-level validation
    parsing_warnings: list[str] = field(default_factory=list)
    
    @property
    def flagged_comments(self) -> list[StudentComment]:
        """Returns comments that have validation flags."""
        return [c for c in self.comments if c.flags]
    
    @property
    def clean_comments(self) -> list[StudentComment]:
        """Returns comments with no validation flags."""
        return [c for c in self.comments if not c.flags]


# =============================================================================
# Parser
# =============================================================================

class DocumentParser:
    """
    Parses teacher comment documents into structured records.
    
    Supports multiple document formats with auto-detection:
    1. Combined header: "LastName, FirstName - Grade" (most common)
    2. Separate header: Name line, then grade line, then comment
    3. Table format: Columns for name, grade, comment
    """
    
    # Expected comment length (from guidelines: 200-250 words)
    MIN_WORD_COUNT = 150   # Allow some buffer below guideline
    MAX_WORD_COUNT = 300   # Allow some buffer above guideline
    TARGET_MIN = 200
    TARGET_MAX = 250
    
    # Combined header pattern: "LastName, FirstName - Grade"
    # Handles: dash variants, inconsistent spacing, various name formats, Unicode
    COMBINED_HEADER_PATTERN = re.compile(
        r"^([\w][\w'\-]+),\s*"               # Last name (Unicode word chars, apostrophe, hyphen)
        r"([\w][\w'\- ]*?)\s*"               # First name (may have middle, spaces)
        r"[-–—]\s*"                          # Dash separator (any variant)
        r"([A-F][+-]?|\d{1,3}(?:\.\d+)?%?)$", # Grade
        re.IGNORECASE | re.UNICODE
    )
    
    # Separate name-only patterns (no grade on line)
    NAME_ONLY_PATTERNS = [
        re.compile(r"^([A-Za-z][A-Za-z'\-]+),\s*([A-Za-z][A-Za-z'\- ]+)$"),  # "Last, First"
        re.compile(r"^([A-Za-z][A-Za-z'\- ]+)\s+([A-Za-z][A-Za-z'\-]+)$"),   # "First Last"
    ]
    
    # Grade-only patterns
    GRADE_PATTERNS = [
        re.compile(r"^Grade[:\s]*([A-F][+-]?)$", re.IGNORECASE),
        re.compile(r"^([A-F][+-]?)$"),
        re.compile(r"^Grade[:\s]*(\d{1,3}(?:\.\d+)?%?)$", re.IGNORECASE),
        re.compile(r"^(\d{1,3}(?:\.\d+)?%?)$"),
    ]
    
    # Section delimiters
    DELIMITER_PATTERNS = [
        re.compile(r"^[-=_]{3,}$"),
        re.compile(r"^\*{3,}$"),
    ]
    
    def __init__(self):
        self._format_detection_cache: dict[str, DocumentFormat] = {}
    
    def parse_docx(self, file_path: Path, document_id: str | None = None) -> TeacherDocument:
        """
        Parse a Word document into a TeacherDocument.
        
        Args:
            file_path: Path to the .docx file
            document_id: Optional ID (generated if not provided)
            
        Returns:
            TeacherDocument with extracted comments
        """
        document_id = document_id or str(uuid.uuid4())
        doc = Document(file_path)
        
        logger.info("parsing_document", path=str(file_path), doc_id=document_id)
        
        # Detect format
        detected_format = self._detect_format(doc)
        logger.info("format_detected", format=detected_format.name)
        
        warnings = []
        
        # Parse based on detected format
        if detected_format == DocumentFormat.TABLE:
            if doc.tables:
                comments = list(self._parse_table_format(doc.tables[0], document_id))
            else:
                warnings.append("Table format detected but no tables found")
                comments = []
        elif detected_format == DocumentFormat.COMBINED_HEADER:
            comments = list(self._parse_combined_header_format(doc, document_id))
        elif detected_format == DocumentFormat.SEPARATE_HEADER:
            comments = list(self._parse_separate_header_format(doc, document_id))
        else:
            warnings.append("Could not detect document format, attempting combined header parse")
            comments = list(self._parse_combined_header_format(doc, document_id))
        
        # Validate all comments
        for comment in comments:
            self._validate_comment(comment)
        
        logger.info(
            "parsing_complete",
            comment_count=len(comments),
            flagged_count=sum(1 for c in comments if c.flags)
        )
        
        return TeacherDocument(
            id=document_id,
            source_path=str(file_path),
            detected_format=detected_format,
            comments=comments,
            parsing_warnings=warnings,
        )
    
    def _detect_format(self, doc: Document) -> DocumentFormat:
        """
        Auto-detect the document format by examining structure.
        """
        # Check for tables first
        if doc.tables:
            # Verify it looks like a comment table
            if len(doc.tables[0].rows) > 1:
                header_text = " ".join(c.text.lower() for c in doc.tables[0].rows[0].cells)
                if any(kw in header_text for kw in ["name", "student", "comment", "grade"]):
                    return DocumentFormat.TABLE
        
        # Examine first several non-empty paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()][:10]
        
        combined_matches = 0
        name_only_matches = 0
        
        for para in paragraphs:
            if self.COMBINED_HEADER_PATTERN.match(para):
                combined_matches += 1
            elif any(p.match(para) for p in self.NAME_ONLY_PATTERNS):
                name_only_matches += 1
        
        if combined_matches >= 1:
            return DocumentFormat.COMBINED_HEADER
        elif name_only_matches >= 1:
            return DocumentFormat.SEPARATE_HEADER
        
        return DocumentFormat.UNKNOWN
    
    def _parse_combined_header_format(
        self, doc: Document, document_id: str
    ) -> Iterator[StudentComment]:
        """
        Parse documents with combined "Name - Grade" headers.
        
        Expected format:
            LastName, FirstName - A+
            
            Comment paragraph text...
            
            NextLastName, NextFirstName - B
            
            Next comment...
        """
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        current_header = None
        current_last = None
        current_first = None
        current_grade = None
        comment_lines = []
        section_index = 0
        
        for para in paragraphs:
            match = self.COMBINED_HEADER_PATTERN.match(para)
            
            if match:
                # Save previous student if exists
                if current_header is not None and comment_lines:
                    section_index += 1
                    yield self._create_comment(
                        document_id=document_id,
                        section_index=section_index,
                        raw_header=current_header,
                        last_name=current_last,
                        first_name=current_first,
                        grade=current_grade,
                        comment_lines=comment_lines,
                    )
                
                # Start new student
                current_header = para
                current_last = match.group(1).strip()
                current_first = match.group(2).strip()
                current_grade = match.group(3).strip()
                comment_lines = []
            else:
                # Accumulate comment text
                comment_lines.append(para)
        
        # Don't forget last student
        if current_header is not None and comment_lines:
            section_index += 1
            yield self._create_comment(
                document_id=document_id,
                section_index=section_index,
                raw_header=current_header,
                last_name=current_last,
                first_name=current_first,
                grade=current_grade,
                comment_lines=comment_lines,
            )
    
    def _parse_separate_header_format(
        self, doc: Document, document_id: str
    ) -> Iterator[StudentComment]:
        """
        Parse documents with name and grade on separate lines.
        """
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        current_header = None
        current_last = None
        current_first = None
        current_grade = None
        comment_lines = []
        section_index = 0
        
        for para in paragraphs:
            # Check for delimiter
            if any(p.match(para) for p in self.DELIMITER_PATTERNS):
                if current_header and comment_lines:
                    section_index += 1
                    yield self._create_comment(
                        document_id=document_id,
                        section_index=section_index,
                        raw_header=current_header,
                        last_name=current_last,
                        first_name=current_first,
                        grade=current_grade,
                        comment_lines=comment_lines,
                    )
                current_header = None
                current_last = None
                current_first = None
                current_grade = None
                comment_lines = []
                continue
            
            # Check for name-only line
            name_match = None
            for pattern in self.NAME_ONLY_PATTERNS:
                name_match = pattern.match(para)
                if name_match:
                    break
            
            if name_match:
                # Save previous if exists
                if current_header and comment_lines:
                    section_index += 1
                    yield self._create_comment(
                        document_id=document_id,
                        section_index=section_index,
                        raw_header=current_header,
                        last_name=current_last,
                        first_name=current_first,
                        grade=current_grade,
                        comment_lines=comment_lines,
                    )
                
                current_header = para
                # Determine if Last, First or First Last
                if "," in para:
                    current_last = name_match.group(1).strip()
                    current_first = name_match.group(2).strip()
                else:
                    current_first = name_match.group(1).strip()
                    current_last = name_match.group(2).strip()
                current_grade = None
                comment_lines = []
                continue
            
            # Check for grade line
            grade_match = None
            for pattern in self.GRADE_PATTERNS:
                grade_match = pattern.match(para)
                if grade_match:
                    break
            
            if grade_match and current_header:
                current_grade = grade_match.group(1).strip()
                continue
            
            # Otherwise it's comment text
            if current_header:
                comment_lines.append(para)
        
        # Last student
        if current_header and comment_lines:
            section_index += 1
            yield self._create_comment(
                document_id=document_id,
                section_index=section_index,
                raw_header=current_header,
                last_name=current_last,
                first_name=current_first,
                grade=current_grade,
                comment_lines=comment_lines,
            )
    
    def _parse_table_format(
        self, table: Table, document_id: str
    ) -> Iterator[StudentComment]:
        """Parse a table-formatted document."""
        header_row = [cell.text.strip().lower() for cell in table.rows[0].cells]
        
        # Find columns
        name_col = self._find_column(header_row, ["name", "student", "student name"])
        grade_col = self._find_column(header_row, ["grade", "mark", "score"])
        comment_col = self._find_column(header_row, ["comment", "comments", "feedback", "notes"])
        
        # Try to find separate first/last columns
        first_col = self._find_column(header_row, ["first", "first name", "given"])
        last_col = self._find_column(header_row, ["last", "last name", "surname", "family"])
        
        if comment_col is None:
            logger.warning("table_missing_comment_column", headers=header_row)
            return
        
        for idx, row in enumerate(table.rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]
            
            # Extract name
            if first_col is not None and last_col is not None:
                first_name = cells[first_col] if first_col < len(cells) else ""
                last_name = cells[last_col] if last_col < len(cells) else ""
                full_name = f"{last_name}, {first_name}"
                raw_header = full_name
            elif name_col is not None:
                raw_header = cells[name_col] if name_col < len(cells) else ""
                last_name, first_name = self._parse_name_string(raw_header)
                full_name = f"{last_name}, {first_name}" if last_name else raw_header
            else:
                continue
            
            grade = cells[grade_col] if grade_col is not None and grade_col < len(cells) else ""
            comment = cells[comment_col] if comment_col < len(cells) else ""
            
            if not comment:
                continue
            
            yield self._create_comment(
                document_id=document_id,
                section_index=idx,
                raw_header=raw_header,
                last_name=last_name,
                first_name=first_name,
                grade=grade,
                comment_lines=[comment],
            )
    
    def _create_comment(
        self,
        document_id: str,
        section_index: int,
        raw_header: str,
        last_name: str,
        first_name: str,
        grade: str | None,
        comment_lines: list[str],
    ) -> StudentComment:
        """Create a StudentComment with computed fields."""
        comment_text = " ".join(comment_lines)
        word_count = len(comment_text.split())
        
        # Find names mentioned in the comment
        names_found = self._find_names_in_text(comment_text, first_name)
        
        return StudentComment(
            id=f"{document_id}-{section_index}",
            document_id=document_id,
            section_index=section_index,
            raw_header=raw_header,
            full_name=f"{last_name}, {first_name}",
            last_name=last_name,
            first_name=first_name,
            grade=grade or "",
            comment_text=comment_text,
            word_count=word_count,
            names_found_in_comment=names_found,
        )
    
    def _validate_comment(self, comment: StudentComment) -> None:
        """
        Validate a comment and add appropriate flags.
        """
        # Empty comment
        if not comment.comment_text.strip():
            comment.flags.append(ValidationFlag.EMPTY_COMMENT)
            return
        
        # Word count checks
        if comment.word_count < self.MIN_WORD_COUNT:
            comment.flags.append(ValidationFlag.SHORT_COMMENT)
            comment.flag_details["word_count"] = (
                f"{comment.word_count} words (expected {self.TARGET_MIN}-{self.TARGET_MAX})"
            )
        elif comment.word_count > self.MAX_WORD_COUNT:
            comment.flags.append(ValidationFlag.LONG_COMMENT)
            comment.flag_details["word_count"] = (
                f"{comment.word_count} words (expected {self.TARGET_MIN}-{self.TARGET_MAX})"
            )
        
        # Grade checks
        if not comment.grade:
            comment.flags.append(ValidationFlag.NO_GRADE)
        elif not re.match(r"^[A-F][+-]?$|^\d{1,3}%?$", comment.grade, re.IGNORECASE):
            comment.flags.append(ValidationFlag.UNUSUAL_GRADE)
            comment.flag_details["grade"] = f"Unexpected grade format: '{comment.grade}'"
        
        # Name usage checks
        first_name_lower = comment.first_name.lower()
        if first_name_lower:
            comment_lower = comment.comment_text.lower()
            
            # Check if first name is used in comment
            if first_name_lower not in comment_lower:
                comment.flags.append(ValidationFlag.NO_FIRST_NAME_USAGE)
                comment.flag_details["name_usage"] = (
                    f"First name '{comment.first_name}' not found in comment"
                )
            
            # Check for other names that appear more frequently
            for name in comment.names_found_in_comment:
                if name.lower() != first_name_lower:
                    name_count = comment_lower.count(name.lower())
                    first_count = comment_lower.count(first_name_lower)
                    if name_count > first_count:
                        comment.flags.append(ValidationFlag.POSSIBLE_WRONG_NAME)
                        comment.flag_details["wrong_name"] = (
                            f"'{name}' appears {name_count}x but expected "
                            f"'{comment.first_name}' appears {first_count}x"
                        )
                        break
    
    def _find_names_in_text(self, text: str, expected_first: str) -> list[str]:
        """
        Find potential first names in the comment text.
        
        Looks for capitalized words that could be names.
        """
        # Simple heuristic: find capitalized words that aren't at sentence start
        # and aren't common words
        common_non_names = {
            "the", "a", "an", "is", "are", "was", "were", "be", "been",
            "have", "has", "had", "do", "does", "did", "will", "would",
            "could", "should", "may", "might", "must", "shall", "can",
            "this", "that", "these", "those", "i", "you", "he", "she",
            "it", "we", "they", "his", "her", "its", "their", "my",
            "your", "our", "after", "before", "during", "while", "if",
            "when", "where", "who", "what", "why", "how", "and", "or",
            "but", "not", "with", "from", "for", "about", "into",
        }
        
        # Find words that look like names
        potential_names = set()
        words = text.split()
        
        for i, word in enumerate(words):
            # Clean punctuation
            clean = re.sub(r"[^\w'-]", "", word)
            if not clean:
                continue
            
            # Check if capitalized and not at sentence start (roughly)
            if clean[0].isupper() and clean.lower() not in common_non_names:
                # Skip if it's the first word or follows sentence-ending punctuation
                if i > 0:
                    prev = words[i-1]
                    if not prev.endswith((".", "!", "?")):
                        potential_names.add(clean)
        
        # Always include expected first name if present
        if expected_first:
            for name in list(potential_names):
                if name.lower() == expected_first.lower():
                    potential_names.add(expected_first)
        
        return list(potential_names)
    
    def _find_column(self, headers: list[str], candidates: list[str]) -> int | None:
        """Find the index of a column matching any candidate name."""
        for idx, header in enumerate(headers):
            if any(c in header for c in candidates):
                return idx
        return None
    
    def _parse_name_string(self, name_str: str) -> tuple[str, str]:
        """Parse a name string into (last_name, first_name)."""
        if "," in name_str:
            parts = name_str.split(",", 1)
            return parts[0].strip(), parts[1].strip()
        else:
            parts = name_str.split()
            if len(parts) >= 2:
                return parts[-1], " ".join(parts[:-1])
            return name_str, ""


# =============================================================================
# Roster Integration
# =============================================================================

class RosterLoader:
    """Loads student rosters and can match against parsed comments."""
    
    @staticmethod
    def from_csv(file_path: Path) -> list[dict]:
        """
        Load roster from CSV file.
        
        Expected columns: student_id, first_name, last_name, preferred_name (optional)
        """
        import csv
        
        roster = []
        with open(file_path, newline='', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                roster.append({
                    'student_id': row.get('student_id', row.get('id', '')),
                    'first_name': row.get('first_name', row.get('first', '')),
                    'last_name': row.get('last_name', row.get('last', '')),
                    'preferred_name': row.get('preferred_name', row.get('nickname', '')),
                })
        
        logger.info("roster_loaded", count=len(roster), source=str(file_path))
        return roster
    
    @staticmethod
    def match_comment_to_roster(
        comment: StudentComment, roster: list[dict]
    ) -> dict | None:
        """
        Find the roster entry matching a comment's student.
        
        Returns the matching roster entry or None.
        """
        comment_last = comment.last_name.lower()
        comment_first = comment.first_name.lower()
        
        for entry in roster:
            roster_last = entry['last_name'].lower()
            roster_first = entry['first_name'].lower()
            roster_pref = entry.get('preferred_name', '').lower()
            
            if roster_last == comment_last:
                if roster_first == comment_first or roster_pref == comment_first:
                    return entry
        
        return None


# =============================================================================
# Convenience Functions
# =============================================================================

def parse_document(file_path: str | Path) -> TeacherDocument:
    """
    Parse a document file into a TeacherDocument.
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        Parsed TeacherDocument with validation flags
    """
    parser = DocumentParser()
    return parser.parse_docx(Path(file_path))


def print_validation_report(doc: TeacherDocument) -> None:
    """Print a human-readable validation report."""
    print(f"\n{'='*60}")
    print(f"Document: {doc.source_path}")
    print(f"Format detected: {doc.detected_format.name}")
    print(f"Total comments: {len(doc.comments)}")
    print(f"Flagged comments: {len(doc.flagged_comments)}")
    print(f"Clean comments: {len(doc.clean_comments)}")
    
    if doc.parsing_warnings:
        print(f"\nParsing Warnings:")
        for w in doc.parsing_warnings:
            print(f"  ⚠ {w}")
    
    if doc.flagged_comments:
        print(f"\nFlagged Comments:")
        for c in doc.flagged_comments:
            print(f"\n  [{c.section_index}] {c.full_name} (Grade: {c.grade})")
            print(f"      Word count: {c.word_count}")
            for flag in c.flags:
                detail = c.flag_details.get(flag.name.lower(), "")
                print(f"      🚩 {flag.name}: {detail}")
    
    print(f"\n{'='*60}\n")


# =============================================================================
# Main (for testing)
# =============================================================================

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        doc = parse_document(sys.argv[1])
        print_validation_report(doc)
        
        print("\nParsed Comments:")
        for c in doc.comments:
            print(f"\n--- {c.full_name} ({c.grade}) ---")
            print(f"First name: {c.first_name}")
            print(f"Word count: {c.word_count}")
            print(f"Flags: {[f.name for f in c.flags]}")
            print(f"Comment preview: {c.comment_text[:100]}...")
    else:
        print("Usage: python stage_0_ingestion_improved.py <document.docx>")
